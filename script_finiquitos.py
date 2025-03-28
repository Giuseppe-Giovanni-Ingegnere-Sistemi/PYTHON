import pandas as pd
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime
from num2words import num2words
from tkinter import Tk, filedialog
import re

# 📌 Función para seleccionar archivos con un cuadro de diálogo
def seleccionar_archivo(tipo):
    root = Tk()
    root.withdraw()  # Oculta la ventana de Tkinter
    if tipo == "excel":
        archivo = filedialog.askopenfilename(title="Selecciona el archivo de Excel", filetypes=[("Archivos de Excel", "*.xlsx")])
    elif tipo == "word":
        archivo = filedialog.askopenfilename(title="Selecciona la plantilla de Word", filetypes=[("Archivos de Word", "*.docx")])
    else:
        archivo = filedialog.askdirectory(title="Selecciona la carpeta donde guardar los documentos")
    return archivo

# Función para formatear valores numéricos correctamente
def formatear_valor(valor):
    if pd.isna(valor):
        return "0.00"
    elif isinstance(valor, (int, float)):
        return f"{valor:,.2f}".replace(",", " ")
    # Si es una fecha, formatearla correctamente en formato AAAA-MM-DD
    elif isinstance(valor, pd.Timestamp) or isinstance(valor, datetime):
        return valor.strftime("%Y-%m-%d")
    return str(valor).strip()

# Función para convertir números a texto en español
def numero_a_texto(numero):
    if pd.isna(numero) or numero == 0:
        return "CERO PESOS 00/100 M.N."
    
    # Separar parte entera y decimal
    partes = str(round(numero, 2)).split('.')
    entero = int(partes[0])
    decimal = int(partes[1]) if len(partes) > 1 else 0
    
    # Convertir a palabras en español
    try:
        texto_entero = num2words(entero, lang='es').upper()
        
        # Formatear el resultado final con paréntesis
        return f"{formatear_valor(numero)} ({texto_entero} PESOS {decimal:02d}/100 M.N.)"
    except:
        # Si hay algún error, devolver un formato básico
        return f"{formatear_valor(numero)} (PESOS {decimal:02d}/100 M.N.)"

# Función para obtener solo el valor numérico formateado
def solo_valor_numerico(numero):
    if pd.isna(numero) or numero == 0:
        return "0.00"
    return formatear_valor(numero)

# Función para obtener el valor con texto para NETO
def valor_neto_con_texto(numero):
    if pd.isna(numero) or numero == 0:
        return "CERO PESOS 00/100 M.N."
    
    # Separar parte entera y decimal
    partes = str(round(numero, 2)).split('.')
    entero = int(partes[0])
    decimal = int(partes[1]) if len(partes) > 1 else 0
    
    # Convertir a palabras en español
    try:
        texto_entero = num2words(entero, lang='es').upper()
        
        # Formato específico para NETO con texto (sin paréntesis)
        return f"{texto_entero} PESOS {decimal:02d}/100 M.N."
    except:
        # Si hay algún error, devolver un formato básico
        return f"PESOS {decimal:02d}/100 M.N."

# Función para reemplazar texto sin perder formato
def reemplazar_texto(doc, buscar, reemplazo, convertir_a_texto=False):
    # Si es un valor numérico y se debe convertir a texto
    if convertir_a_texto and isinstance(reemplazo, (int, float)):
        texto_completo = numero_a_texto(reemplazo)
    else:
        texto_completo = formatear_valor(reemplazo)
    
    for parrafo in doc.paragraphs:
        if buscar in parrafo.text:
            for run in parrafo.runs:
                if buscar in run.text:
                    run.text = run.text.replace(buscar, texto_completo)
                    run.bold = True
                    run.font.name = "Verdana"
                    run.font.size = Pt(11)
    
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    if buscar in parrafo.text:
                        for run in parrafo.runs:
                            if buscar in run.text:
                                run.text = run.text.replace(buscar, texto_completo)
                                run.bold = True
                                run.font.name = "Verdana"
                                run.font.size = Pt(11)

# Función especial para manejar las diferentes instancias de NETO
def reemplazar_neto(doc, valor_neto):
    if pd.isna(valor_neto):
        valor_neto = 0
    
    # Asegurarse de que sea un número
    if isinstance(valor_neto, str):
        try:
            valor_neto = float(valor_neto.replace(',', '').replace(' ', ''))
        except:
            valor_neto = 0
    
    # Obtener los valores formateados
    valor_numerico = solo_valor_numerico(valor_neto)
    texto_completo = valor_neto_con_texto(valor_neto)
    
    # Patrones para identificar los diferentes contextos
    patrones_solo_numero = [
        r"BUENO POR: \$ «NETO»",
        r"Total neto a recibir\s+\$ «NETO»"
    ]
    
    patrones_con_texto = [
        r"Recibí la cantidad de \$ «NETO»",
        r"se realizará por \$ «NETO»"
    ]
    
    # Reemplazar en párrafos
    for parrafo in doc.paragraphs:
        texto_original = parrafo.text
        texto_modificado = texto_original
        
        # Verificar si es un patrón que solo necesita el número
        for patron in patrones_solo_numero:
            if re.search(patron, texto_original):
                texto_modificado = texto_modificado.replace("«NETO»", valor_numerico)
                break
        
        # Verificar si es un patrón que necesita número y texto
        for patron in patrones_con_texto:
            if re.search(patron, texto_original):
                texto_modificado = texto_modificado.replace("«NETO»", f"{valor_numerico} ({texto_completo})")
                break
        
        # Si no coincide con ningún patrón específico, usar solo el número
        if "«NETO»" in texto_modificado:
            texto_modificado = texto_modificado.replace("«NETO»", valor_numerico)
        
        # Si hubo cambios, actualizar el texto manteniendo el formato
        if texto_modificado != texto_original:
            for run in parrafo.runs:
                if "«NETO»" in run.text:
                    # Determinar qué reemplazo usar
                    for patron in patrones_solo_numero:
                        if re.search(patron, texto_original):
                            run.text = run.text.replace("«NETO»", valor_numerico)
                            break
                    
                    for patron in patrones_con_texto:
                        if re.search(patron, texto_original):
                            run.text = run.text.replace("«NETO»", f"{valor_numerico} ({texto_completo})")
                            break
                    
                    # Si no coincide con ningún patrón específico, usar solo el número
                    if "«NETO»" in run.text:
                        run.text = run.text.replace("«NETO»", valor_numerico)
                    
                    run.bold = True
                    run.font.name = "Verdana"
                    run.font.size = Pt(11)
    
    # Reemplazar en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    texto_original = parrafo.text
                    
                    # Para tablas, generalmente solo usamos el valor numérico
                    if "«NETO»" in texto_original:
                        for run in parrafo.runs:
                            if "«NETO»" in run.text:
                                run.text = run.text.replace("«NETO»", valor_numerico)
                                run.bold = True
                                run.font.name = "Verdana"
                                run.font.size = Pt(11)

# Función especial para manejar las diferentes instancias de Salario_por_día
def reemplazar_salario(doc, valor_salario):
    if pd.isna(valor_salario):
        valor_salario = 0
    
    # Asegurarse de que sea un número
    if isinstance(valor_salario, str):
        try:
            valor_salario = float(valor_salario.replace(',', '').replace(' ', ''))
        except:
            valor_salario = 0
    
    # Obtener los valores formateados
    valor_numerico = solo_valor_numerico(valor_salario)
    
    # Convertir a palabras en español
    partes = str(round(valor_salario, 2)).split('.')
    entero = int(partes[0])
    decimal = int(partes[1]) if len(partes) > 1 else 0
    
    try:
        texto_entero = num2words(entero, lang='es').upper()
        texto_completo = f"{texto_entero} PESOS {decimal:02d}/100 M.N."
    except:
        texto_completo = f"PESOS {decimal:02d}/100 M.N."
    
    # Patrones para identificar los diferentes contextos
    patrones_con_texto = [
        r"SALARIO DIARIO: \$ «Salario_por_día»",
        r"salario diario por la cantidad de \$ «Salario_por_día»"
    ]
    
    # Reemplazar en párrafos
    for parrafo in doc.paragraphs:
        texto_original = parrafo.text
        
        # Verificar si es un patrón que necesita número y texto
        necesita_texto = False
        for patron in patrones_con_texto:
            if re.search(patron, texto_original, re.IGNORECASE) or "salario diario por la cantidad de $ «Salario_por_día»" in texto_original:
                necesita_texto = True
                break
        
        if necesita_texto:
            for run in parrafo.runs:
                if "«Salario_por_día»" in run.text:
                    run.text = run.text.replace("«Salario_por_día»", f"{valor_numerico} ({texto_completo})")
                    run.bold = True
                    run.font.name = "Verdana"
                    run.font.size = Pt(11)
        # Para otros casos donde solo se necesita el número
        elif "«Salario_por_día»" in texto_original:
            for run in parrafo.runs:
                if "«Salario_por_día»" in run.text:
                    run.text = run.text.replace("«Salario_por_día»", valor_numerico)
                    run.bold = True
                    run.font.name = "Verdana"
                    run.font.size = Pt(11)
    
    # Reemplazar en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    texto_original = parrafo.text
                    
                    # Para tablas, generalmente solo usamos el valor numérico
                    if "«Salario_por_día»" in texto_original:
                        for run in parrafo.runs:
                            if "«Salario_por_día»" in run.text:
                                run.text = run.text.replace("«Salario_por_día»", valor_numerico)
                                run.bold = True
                                run.font.name = "Verdana"
                                run.font.size = Pt(11)

# 📌 Pedir al usuario que seleccione los archivos
ruta_excel = seleccionar_archivo("excel")
ruta_plantilla = seleccionar_archivo("word")
ruta_guardar = seleccionar_archivo("carpeta")

# 📌 Si el usuario no selecciona nada, salir del programa
if not ruta_excel or not ruta_plantilla or not ruta_guardar:
    print("⚠️ No seleccionaste todos los archivos. Saliendo del programa...")
    exit()

# 📌 Leer el archivo Excel correctamente
df = pd.read_excel(ruta_excel, sheet_name="CALCULO", header=5)
df.columns = df.columns.str.strip()  # Eliminar espacios en los nombres de columnas

# 📌 Crear la carpeta de guardado si no existe
if not os.path.exists(ruta_guardar):
    os.makedirs(ruta_guardar)

# 📌 Iterar sobre todas las filas del Excel para generar un documento por cada persona
for index, fila in df.iterrows():
    if pd.isna(fila.get("Nombre completo", "")):
        continue

    # 📌 Abrir la plantilla de Word
    doc = Document(ruta_plantilla)

    # 📌 Reemplazar los marcadores en el documento
    reemplazar_texto(doc, "«Nombre_completo»", fila.get("Nombre completo", ""))
    reemplazar_texto(doc, "«Puesto»", fila.get("Puesto", ""))
    
    # Usar la función especial para Salario_por_día
    salario_dia = fila.get("Salario por día", 0)
    if isinstance(salario_dia, str):
        try:
            salario_dia = float(salario_dia.replace(',', '').replace(' ', ''))
        except:
            salario_dia = 0
    reemplazar_salario(doc, salario_dia)
    
    reemplazar_texto(doc, "«Fecha_de_alta»", fila.get("Fecha de alta", ""))
    reemplazar_texto(doc, "«Fecha_de_baja»", fila.get("Fecha de baja", ""))
    reemplazar_texto(doc, "«SUELDO»", fila.get("SUELDO", ""))
    reemplazar_texto(doc, "«IMPORTE_AGUINALDO»", fila.get("IMPORTE AGUINALDO", ""))
    reemplazar_texto(doc, "«IMPORTE_VACACIONES»", fila.get("IMPORTE VACACIONES", ""))
    reemplazar_texto(doc, "«GRATIFICACION»", fila.get("GRATIFICACION", ""))
    reemplazar_texto(doc, "«IMPORTE_PRIMA_VACACIONAL»", fila.get("IMPORTE PRIMA VACACIONAL", ""))
    reemplazar_texto(doc, "«Total_de_Percepciones»", fila.get("TOTAL PERCEPCIONES", ""))
    reemplazar_texto(doc, "«TOTAL_ISR»", fila.get("ISR MENSUAL", ""))
    reemplazar_texto(doc, "«IMSS»", fila.get("IMSS", ""))
    reemplazar_texto(doc, "«TOTAL_DEDUCCIONES»", fila.get("TOTAL DEDUCCIONES", ""))
    
    # Usar la función especial para NETO
    reemplazar_neto(doc, fila.get("NETO", 0))
    
    reemplazar_texto(doc, "«Banco»", fila.get("Banco", ""))
    reemplazar_texto(doc, "«CUENTA»", fila.get("cuenta", ""))

    # 📌 Guardar el documento con el nombre del trabajador
    nombre_archivo = os.path.join(ruta_guardar, f"{fila.get('Nombre completo', 'SinNombre').replace('/', '-').replace(':', '-')}.docx")
    doc.save(nombre_archivo)
    print(f"✅ Documento generado: {nombre_archivo}")

print("✅ TODOS los documentos han sido generados correctamente.")