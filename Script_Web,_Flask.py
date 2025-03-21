from flask import Flask, render_template, request, send_from_directory
import pandas as pd
from docx import Document
from docx.shared import Pt
import os
from werkzeug.utils import secure_filename

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER

def formatear_valor(valor):
    if pd.isna(valor):
        return "0.00"
    elif isinstance(valor, (int, float)):
        return f"{valor:,.2f}".replace(",", " ")
    return str(valor).strip()

def reemplazar_texto(doc, buscar, reemplazo):
    reemplazo = formatear_valor(reemplazo)
    for parrafo in doc.paragraphs:
        if buscar in parrafo.text:
            for run in parrafo.runs:
                if buscar in run.text:
                    run.text = run.text.replace(buscar, reemplazo)
                    run.bold = True
                    run.font.name = "Verdana"
                    run.font.size = Pt(11)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    if buscar in parrafo.text:
                        for run in parrafo.runs:
                            if buscar in run.text:
                                run.text = run.text.replace(buscar, reemplazo)
                                run.bold = True
                                run.font.name = "Verdana"
                                run.font.size = Pt(11)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        excel_file = request.files['excel']
        word_file = request.files['word']
        
        if not excel_file or not word_file:
            return render_template('index.html', mensaje="⚠️ Debes subir ambos archivos.")
        
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(excel_file.filename))
        word_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(word_file.filename))
        
        excel_file.save(excel_path)
        word_file.save(word_path)
        
        df = pd.read_excel(excel_path, sheet_name="CALCULO", header=5)
        df.columns = df.columns.str.strip()
        
        for index, fila in df.iterrows():
            if pd.isna(fila.get("Nombre completo", "")):
                continue
            
            doc = Document(word_path)
            placeholders = [
                ("«Nombre_completo»", fila.get("Nombre completo", "")),
                ("«Puesto»", fila.get("Puesto", "")),
                ("«Salario_por_día»", fila.get("Salario por día", "")),
                ("«Fecha_de_alta»", fila.get("Fecha de alta", "")),
                ("«Fecha_de_baja»", fila.get("Fecha de baja", "")),
                ("«SUELDO»", fila.get("SUELDO", "")),
                ("«IMPORTE_AGUINALDO»", fila.get("IMPORTE AGUINALDO", "")),
                ("«IMPORTE_VACACIONES»", fila.get("IMPORTE VACACIONES", "")),
                ("«GRATIFICACION»", fila.get("GRATIFICACION", "")),
                ("«IMPORTE_PRIMA_VACACIONAL»", fila.get("IMPORTE PRIMA VACACIONAL", "")),
                ("«Total_de_Percepciones»", fila.get("TOTAL PERCEPCIONES", "")),
                ("«TOTAL_ISR»", fila.get("ISR MENSUAL", "")),
                ("«IMSS»", fila.get("IMSS", "")),
                ("«TOTAL_DEDUCCIONES»", fila.get("TOTAL DEDUCCIONES", "")),
                ("«NETO»", fila.get("NETO", "")),
                ("«Banco»", fila.get("Banco", "")),
                ("«CUENTA»", fila.get("cuenta", ""))
            ]
            
            for placeholder, value in placeholders:
                reemplazar_texto(doc, placeholder, value)
            
            nombre_archivo = os.path.join(
                app.config['PROCESSED_FOLDER'],
                f"{fila.get('Nombre completo', 'SinNombre').replace('/', '-').replace(':', '-')}.docx"
            )
            doc.save(nombre_archivo)
        
        return render_template('index.html', mensaje="✅ Archivos procesados con éxito. Descarga en la carpeta de procesados.")
    
    return render_template('index.html', mensaje="")

@app.route('/descargas')
def descargas():
    archivos = os.listdir(PROCESSED_FOLDER)
    return render_template('descargas.html', archivos=archivos)

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(PROCESSED_FOLDER, filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)